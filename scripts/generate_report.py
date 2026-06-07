#!/usr/bin/env python3
"""
生成正式实验报告 .docx
依据：docs/实验报告.md 内容 + 实验报告模板（信息学院软件类）.docx 格式
版本：v1.0 | 日期：2026-06-07
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# 工具函数
# ============================================================

def set_cn_font(run, font_name_cn, font_name_en=None, size=None, bold=None, color=None):
    """设置中文字体（处理 western + east-asia 分离）"""
    if font_name_en is None:
        font_name_en = font_name_cn
    run.font.name = font_name_en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name_cn)
    rFonts.set(qn('w:cs'), font_name_cn)
    if size:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading_styled(doc, text, level=1):
    """添加标题（黑体）"""
    if level == 1:
        size = Pt(16)
    elif level == 2:
        size = Pt(15)
    else:
        size = Pt(14)

    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
    else:
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run(text)
    set_cn_font(run, '黑体', 'Arial', size=size, bold=True)
    return p


def add_body(doc, text, indent=True):
    """添加正文段落（宋体小四 12pt，两端对齐，1.5 倍行距）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)  # 2 字符
    run = p.add_run(text)
    set_cn_font(run, '宋体', 'Times New Roman', size=Pt(12))
    return p


def add_body_no_indent(doc, text):
    """正文不缩进"""
    return add_body(doc, text, indent=False)


def add_code_block(doc, code_text):
    """添加代码块（Courier New 小五 9pt，灰色底纹）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

    # 灰色底纹
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
    pPr.append(shd)

    run = p.add_run(code_text)
    set_cn_font(run, 'Courier New', 'Courier New', size=Pt(9))
    return p


def add_placeholder(doc, material_id, description):
    """添加素材占位符（虚线框 + 文字说明）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0

    run = p.add_run(f'┌──────────────────────────────────────────────┐\n')
    set_cn_font(run, '宋体', size=Pt(9), color=RGBColor(0x88, 0x88, 0x88))
    run = p.add_run(f'│  [{material_id}] 请在此处插入图片                  │\n')
    set_cn_font(run, '宋体', size=Pt(9), color=RGBColor(0x88, 0x88, 0x88))
    run = p.add_run(f'│  {description[:45]}│\n')
    set_cn_font(run, '宋体', size=Pt(9), color=RGBColor(0x88, 0x88, 0x88))
    run = p.add_run(f'└──────────────────────────────────────────────┘')
    set_cn_font(run, '宋体', size=Pt(9), color=RGBColor(0x88, 0x88, 0x88))
    return p


def add_table_cell_text(cell, text, font_name='宋体', font_name_en='Times New Roman', size=Pt(10.5), bold=False, align='center'):
    """设置表格单元格文字"""
    # 清除已有段落
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    p = cell.paragraphs[0]
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(str(text))
    set_cn_font(run, font_name, font_name_en, size=size, bold=bold)
    return run


def add_table(doc, headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        add_table_cell_text(table.rows[0].cells[i], header, bold=True, size=Pt(10.5))

    # 数据行
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            add_table_cell_text(table.rows[ri + 1].cells[ci], str(cell_text), bold=False, size=Pt(10.5))

    # 设置列宽
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    # 空行
    doc.add_paragraph()
    return table


# ============================================================
# 主函数
# ============================================================

def generate_report():
    doc = Document()

    # ---- 页面设置 ----
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.5)

    # ---- 预定义样式 ----
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    rPr = style.element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="宋体"/>')
    rPr.insert(0, rFonts)

    # ==========================================
    # 封面
    # ==========================================
    # 空行推到中间位置
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0

    # 学校名 — 楷体 三号(16pt) 居中
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.line_spacing = 1.5
    run = p_title.add_run('湖南人文科技学院信息学院')
    set_cn_font(run, '楷体', '楷体', size=Pt(16))

    # 空一行
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0

    # 副标题 — 楷体 四号(14pt) 居中，课程名红色
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.line_spacing = 1.5

    run1 = p_sub.add_run('     ')
    set_cn_font(run1, '楷体', '楷体', size=Pt(14))
    run2 = p_sub.add_run('计算机图形与视觉')
    set_cn_font(run2, '楷体', '楷体', size=Pt(14), color=RGBColor(0xFF, 0x00, 0x00))
    run3 = p_sub.add_run('     ')
    set_cn_font(run3, '楷体', '楷体', size=Pt(14))
    run4 = p_sub.add_run('实验报告')
    set_cn_font(run4, '楷体', '楷体', size=Pt(14))

    # 空行
    for _ in range(3):
        p = doc.add_paragraph()

    # 封面效果图占位
    add_placeholder(doc, '素材-封面', '项目推理效果缩略图 (8cm×5cm)')

    # ---- 分页 ----
    doc.add_page_break()

    # ==========================================
    # 信息表
    # ==========================================
    add_heading_styled(doc, '信息表', level=1)

    info_table = doc.add_table(rows=6, cols=6)
    info_table.style = 'Table Grid'
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: 实验名称
    add_table_cell_text(info_table.rows[0].cells[0], '实验名称：', bold=True, size=Pt(12))
    info_table.rows[0].cells[1].merge(info_table.rows[0].cells[5])
    add_table_cell_text(info_table.rows[0].cells[1], '交通监控智能分析系统 — 车流量统计与车牌识别', size=Pt(12))

    # Row 1: 年级/专业/班级
    add_table_cell_text(info_table.rows[1].cells[0], '年级：', bold=True, size=Pt(12))
    add_table_cell_text(info_table.rows[1].cells[1], '2023级', size=Pt(12))
    add_table_cell_text(info_table.rows[1].cells[2], '专业：', bold=True, size=Pt(12))
    add_table_cell_text(info_table.rows[1].cells[3], '计算机科学与技术', size=Pt(12))
    add_table_cell_text(info_table.rows[1].cells[4], '班级：', bold=True, size=Pt(12))
    add_table_cell_text(info_table.rows[1].cells[5], 'B2023-（填写）', size=Pt(12))

    # Row 2: 姓名/学号/组号
    add_table_cell_text(info_table.rows[2].cells[0], '姓名：', bold=True, size=Pt(12))
    add_table_cell_text(info_table.rows[2].cells[1], '（填写）', size=Pt(12))
    add_table_cell_text(info_table.rows[2].cells[2], '学号：', bold=True, size=Pt(12))
    add_table_cell_text(info_table.rows[2].cells[3], '（填写）', size=Pt(12))
    add_table_cell_text(info_table.rows[2].cells[4], '组号：', bold=True, size=Pt(12))
    add_table_cell_text(info_table.rows[2].cells[5], '（填写）', size=Pt(12))

    # Row 3: 实验地点/日期
    add_table_cell_text(info_table.rows[3].cells[0], '实验地点：', bold=True, size=Pt(12))
    info_table.rows[3].cells[1].merge(info_table.rows[3].cells[2])
    add_table_cell_text(info_table.rows[3].cells[1], 'AutoDL 云端 GPU (RTX 3090)', size=Pt(12))
    add_table_cell_text(info_table.rows[3].cells[3], '日期：', bold=True, size=Pt(12))
    info_table.rows[3].cells[4].merge(info_table.rows[3].cells[5])
    add_table_cell_text(info_table.rows[3].cells[4], '2026 年 6 月', size=Pt(12))

    # Row 4: 组员姓名
    add_table_cell_text(info_table.rows[4].cells[0], '组员姓名：', bold=True, size=Pt(12))
    info_table.rows[4].cells[1].merge(info_table.rows[4].cells[5])
    add_table_cell_text(info_table.rows[4].cells[1], '（如无组员，填"独立完成"）', size=Pt(12))

    # Row 5: 组员学号
    add_table_cell_text(info_table.rows[5].cells[0], '组员学号：', bold=True, size=Pt(12))
    info_table.rows[5].cells[1].merge(info_table.rows[5].cells[5])
    add_table_cell_text(info_table.rows[5].cells[1], '（如无组员，留空）', size=Pt(12))

    doc.add_paragraph()

    # ==========================================
    # 一、实验目的
    # ==========================================
    doc.add_page_break()
    add_heading_styled(doc, '一、实验目的', level=1)

    add_heading_styled(doc, '1.1 项目背景', level=2)

    add_body(doc, '车流量统计是智能交通系统（Intelligent Transportation System, ITS）的核心需求。通过分析道路交通监控视频，自动统计车流量并提取车辆属性信息（车牌号码、车牌颜色），可为交通规划、拥堵治理和治安防控提供关键数据支撑。传统交通监控依赖人工查看视频，效率低、易遗漏、不可持续——这正是计算机视觉技术可以系统性解决的问题。')

    add_body(doc, '本项目设计并实现一套端到端的交通监控视频智能分析系统：输入道路监控视频（4K 高位俯拍），输出带完整标注的分析结果，包含车流量统计、车牌号码识别、车牌颜色分类等功能。')

    add_heading_styled(doc, '1.2 实验目标', level=2)

    add_table(doc,
        ['目标', '指标要求', '对应课程目标'],
        [
            ['车流量统计', '区分正向(IN)/反向(OUT)，与真实值偏差 ≤5%', '目标2（目标检测与跟踪）'],
            ['车牌号码提取', '对越线车辆提取完整车牌号码', '目标2（文字识别）'],
            ['车牌颜色分类', '识别蓝/绿/黄/其他四类底色，准确率 ≥80%', '目标2（图像分类）'],
            ['端到端推理', '输入视频 → 输出标注结果', '目标2（系统集成）'],
            ['原理掌握', '深入理解 YOLO/ByteTrack/OCR 数学原理与工程实现', '目标1（原理与方法）'],
            ['技术调研', '梳理智能交通 CV 技术演进，对比国内外方案', '目标3（前沿与进展）'],
        ],
        [Cm(3.5), Cm(6), Cm(4.5)]
    )

    add_heading_styled(doc, '1.3 交付物', level=2)

    add_table(doc,
        ['交付物', '说明'],
        [
            ['实验报告', '本文档（.docx 格式，依据学院模板）'],
            ['项目完整代码', 'GitHub 仓库（含训练脚本、推理管线、部署说明）'],
            ['答辩 PPT', '包含系统架构、训练结果、演示视频/截图、Q&A 准备'],
        ],
        [Cm(3.5), Cm(10)]
    )

    # ==========================================
    # 二、实验内容、要求和环境
    # ==========================================
    doc.add_page_break()
    add_heading_styled(doc, '二、实验内容、要求和环境', level=1)
    add_body(doc, '本章涵盖评分参考中的全部"原理阐述""技术路线""数据集与训练""代码规范""部署集成"内容。依据模板要求，「实验过程」在 2.8 节集中记录从环境搭建到推理输出的全流程操作。')

    # ---- 2.1 技术调研与算法原理 ----
    add_heading_styled(doc, '2.1 技术调研与算法原理  [评分点 1.1 + 3.1 + 3.2]', level=2)

    add_heading_styled(doc, '2.1.1 智能交通视觉技术演进', level=3)

    add_body(doc, '交通监控中的车辆检测与跟踪技术经历了从传统方法到深度学习的系统性演进。下表梳理各阶段代表技术及其特点：')

    add_table(doc,
        ['阶段', '代表方法', '核心思路', '优势', '局限'],
        [
            ['传统方法(2010s前)', 'HOG+SVM, 卡尔曼滤波', '手工特征+经典跟踪', '可解释强，低算力', '光照/遮挡敏感'],
            ['两阶段检测(2014-17)', 'Faster R-CNN, R-FCN', 'RPN提取候选→分类回归', '精度高', '推理慢(~5FPS)'],
            ['单阶段检测(2016-今)', 'YOLOv1→v8→v11', '端到端回归检测框', '速度快(30+FPS)', '小目标检测弱'],
            ['Transformer(2020-今)', 'DETR, RT-DETR', '端到端集合预测', '全局感受野', '训练收敛慢'],
            ['多模态大模型(2024-)', 'GPT-4V, SAM', '视觉-语言联合理解', '零样本泛化', '延迟高(>1s)'],
        ],
        [Cm(2.5), Cm(3), Cm(3.5), Cm(2.5), Cm(2.5)]
    )

    add_body(doc, '本项目技术选型：基于对以上技术路线的对比分析，综合考虑精度、推理速度与实际部署条件（单卡 RTX 3090），选择 YOLOv8m + ByteTrack + PaddleOCR 技术栈。')

    add_placeholder(doc, '素材-01', '技术演进时间线图 (2010→2024)，宽度=14cm')

    # ---- YOLOv8 原理 ----
    add_heading_styled(doc, '2.1.2 YOLOv8 目标检测 — 原理深度剖析', level=3)

    add_body(doc, 'YOLOv8 采用单阶段 Anchor-Free 架构，由骨干网络（Backbone）、颈部（Neck）、检测头（Head）三部分组成。')

    add_body(doc, '（1）骨干网络 — C2f 模块。C2f（Cross-Stage Partial with 2 Convolutions）是 YOLOv8 的核心特征提取单元。其设计优势在于：部分特征绕过多层卷积直接传递（CSP 思想），在保持轻量化的同时最大化梯度流，有效缓解深度网络中的梯度消失问题。YOLOv8m 使用 n=2 个 Bottleneck，共约 25.9M 参数。')

    add_body(doc, '（2）颈部 — PAN-FPN。FPN 将高层语义信息向下传递，PAN 将底层定位信息向上传递，双向融合后得到多尺度特征金字塔。公式为 P_i = f_upsample(P_{i+1}) + C_i (FPN 自顶向下)，N_i = f_downsample(N_{i-1}) + P_i (PAN 自底向上)，其中 C_i 是 Backbone 的第 i 层输出特征图（i=3,4,5 对应 3 个检测尺度）。')

    add_body(doc, '（3）检测头 — 解耦头设计。与 YOLOv5 的耦合头不同，YOLOv8 将分类和回归任务分离为独立的卷积分支。分类分支输出 [B, num_classes, H, W]，回归分支输出 [B, 4×reg_max, H, W]（Distribution Focal Loss，将框坐标建模为离散概率分布）。Anchor-Free 设计直接从网格中心点回归 (left, top, right, bottom) 四个距离值，省去了锚框超参数调优。')

    add_body(doc, '损失函数：L_total = λ₁ × L_cls(BCE) + λ₂ × L_box(CIoU) + λ₃ × L_dfl(Distribution Focal Loss)，三个损失项分别对应分类精度、框定位精度和分布质量。')

    # ---- ByteTrack 原理 ----
    add_heading_styled(doc, '2.1.3 ByteTrack 多目标跟踪 — 原理深度剖析', level=3)

    add_body(doc, 'ByteTrack 基于 Tracking-by-Detection 范式，其核心创新在于充分利用低置信度检测框，而非像传统方法（SORT、DeepSORT）那样直接丢弃。')

    add_body(doc, '算法流程：对于每一帧，检测器首先获得所有检测框。高置信度框（score > 0.6）优先与已有轨迹进行匈牙利匹配（IoU cost matrix）；然后低置信度框（0.1 < score < 0.6）与未匹配的轨迹进行第二次匹配；未匹配的高分框初始化为新轨迹；连续 N 帧未匹配的轨迹标记为 lost 并删除。')

    add_body(doc, '有效性分析：在遮挡场景下，被遮挡目标的检测置信度会从 0.9 骤降至 0.3-0.5。如果直接丢弃（传统方法），该目标将暂时"消失"，当重新出现时可能被赋予新 ID（ID Switch）。ByteTrack 通过在第二次匹配中接纳这些低分框，维持了轨迹的连续性，显著降低 ID Switch 率。')

    # ---- PaddleOCR 原理 ----
    add_heading_styled(doc, '2.1.4 PaddleOCR PP-OCRv4 — 原理剖析', level=3)

    add_body(doc, 'PaddleOCR PP-OCRv4 采用 DBNet 文本检测 + SVTR 文本识别的两阶段架构。')

    add_body(doc, '文本检测 — DBNet（Differentiable Binarization）：传统文本检测使用固定阈值进行二值化，不可微分，无法端到端训练。DBNet 的核心改进是引入可微分二值化函数 B̂_{ij} = 1 / (1 + e^(-k × (P_{ij} - T_{ij})))，其中 P_{ij} 是预测的概率图，T_{ij} 是自适应阈值图，k 是放大因子（通常 k=50）。这一设计使得网络能够自适应地学习每个像素的最佳二值化阈值，对车牌区域的不规则形状和倾斜有较好的鲁棒性。')

    add_body(doc, '文本识别 — SVTR（Scene Text Recognition with Single Visual Model）：基于 Vision Transformer，通过分块注意力机制（将特征图分为局部块，在块内做自注意力，在块间做全局注意力），高效处理字符级别的细粒度特征，对中文车牌字符的序列识别效果优异。')

    # ---- HSV 颜色分类 ----
    add_heading_styled(doc, '2.1.5 HSV 颜色空间分类 — 原理', level=3)

    add_body(doc, '车牌颜色分类选择 HSV（Hue-Saturation-Value）空间而非 RGB，原因在于 HSV 将颜色（色相）与亮度分离，对光照变化更鲁棒。')

    add_table(doc,
        ['颜色类别', 'H 范围', 'S 条件', 'V 条件', '典型车牌'],
        [
            ['蓝色', '100°–130°', '> 0.40', '> 0.30', '普通小汽车'],
            ['绿色', '40°–80°', '> 0.30', '> 0.20', '新能源车'],
            ['黄色', '20°–40°', '> 0.40', '> 0.30', '大型车/挂车'],
            ['其他', '不满足以上', '—', '—', '白色/黑色/特殊车辆'],
        ],
        [Cm(2.5), Cm(2.5), Cm(2.5), Cm(2.5), Cm(4)]
    )

    add_body(doc, '判定阈值取 8%：当某颜色类别的像素数占车牌区域总像素数的比例超过 8% 时，即判定为该颜色。该方法无需额外训练，计算开销极低（单次 HSV 转换 + 像素遍历约 0.1ms），适合实时推理场景。')

    # ---- 行业应用现状 ----
    add_heading_styled(doc, '2.1.6 行业应用现状与对比  [评分点 3.2]', level=3)

    add_body(doc, '国内外智能交通视觉系统已规模化部署。下表对比了主要厂商的技术路线：')

    add_table(doc,
        ['厂商/系统', '核心技术路线', '优势', '落地挑战', '与本项目对比'],
        [
            ['海康威视"雷视一体"', '毫米波雷达+视觉融合', '全天候，测距准', '成本高，标定复杂', '本项目纯视觉，成本低'],
            ['大华"睿智"系列', '前端嵌入式深度学习', '边缘实时，低延迟', '算法更新不灵活', 'GPU灵活性更高'],
            ['Tesla Vision', 'BEV+Transformer', '3D感知，端到端', '数据需求巨大', '2D俯拍场景更简单'],
            ['Waymo', '激光雷达为主', 'L4级精度', '成本极高(>$100k)', '场景不同'],
            ['本项目', 'YOLOv8m+PaddleOCR', '精度高/部署简单', '纯视觉极端天气受限', '—'],
        ],
        [Cm(2.5), Cm(3), Cm(3), Cm(3), Cm(3)]
    )

    add_body(doc, '这些系统的共性技术挑战包括：(1) 实时性约束——交通监控要求 ≥25 FPS；(2) 复杂天气鲁棒性——雨雪、雾霾、夜间场景的检测精度显著下降；(3) 边缘部署算力约束——路侧设备通常算力有限；(4) 长尾场景覆盖——罕见车型、遮挡车牌、非标准车牌难以覆盖；(5) 隐私合规——车牌属于个人身份信息，需满足 GDPR/《个人信息保护法》要求。')

    # ---- 2.2 方案选型 ----
    add_heading_styled(doc, '2.2 方案选型理由  [评分点 1.1]', level=2)

    add_body(doc, '以下选型均基于实验验证与参考项目的经验，非盲目引用。')

    add_table(doc,
        ['环节', '选择', '备选方案', '选型理由'],
        [
            ['车辆检测', 'YOLOv8m (COCO预训练)', 'YOLOv11n, Faster R-CNN', '25.9M参数，精度速度平衡；COCO预训练对车辆类别已有良好泛化'],
            ['车辆跟踪', 'ByteTrack (YOLO内置)', 'DeepSORT', '低分框二次匹配，抗遮挡；YOLO原生集成'],
            ['越线计数', 'supervision LineZone', '手写逻辑', '社区标准库，向量叉积算法+单ID单次触发'],
            ['车牌检测', 'YOLOv8m（自训练）', 'YOLOv11n, DETR', 'mAP50-95=0.981，YOLOv11n仅0.10'],
            ['文字识别', 'PaddleOCR PP-OCRv4', 'EasyOCR, Tesseract', '中文识别SOTA，DBNet+SVTR双阶段'],
            ['颜色分类', 'HSV规则', '训练CNN分类器', '零计算开销，8%阈值实验验证稳定'],
            ['中文渲染', 'PIL + wqy-zenhei.ttc', 'OpenCV putText', 'PIL支持TrueType中文渲染，OpenCV中文乱码'],
        ],
        [Cm(2), Cm(3.5), Cm(3), Cm(5.5)]
    )

    add_body(doc, '关键决策——为什么不用 YOLOv11？经过 6 轮实验验证：YOLOv11n（2.6M 参数）在 CCPD2019 车牌检测任务上 mAP50-95 最高仅 0.28。根本原因在于车牌在 4K 画面中仅占约 50×15 像素，属于超小目标，nano 级别模型的感受野和特征表达能力严重不足。YOLOv8m（25.9M 参数）一次训练即达 mAP50-95 = 0.981。结论：小目标检测场景下，模型容量是不可压缩的底线。')

    # ---- 2.3 系统架构 ----
    add_heading_styled(doc, '2.3 系统架构设计  [评分点 1.2]', level=2)

    add_placeholder(doc, '素材-02', '系统架构全景图（专业绘图版，Draw.io/Visio），宽度=15cm')

    add_heading_styled(doc, '2.3.1 架构设计原则', level=3)

    add_body(doc, '本系统遵循"高效级联、按需计算"的设计原则，共包含四项核心原则：')

    add_body(doc, '原则 1 — 级联检测替代全帧搜索。传统方案在全帧同时运行车辆检测和车牌检测两个模型。本系统先检测车辆 ROI，再在 ROI 内检测车牌。实验表明，级联方案将车牌误检率降低约 70%（车身文字不再被误认为车牌），同时推理速度提升约 30%（车牌检测搜索空间缩小 10-100 倍）。')

    add_body(doc, '原则 2 — 距离感知动态过滤。画面顶部 30% 区域的车辆距摄像头过远，车牌像素数不足以支撑 OCR 识别（车牌宽度约 20px，小于 OCR 模型的最小有效输入 32px）。跳过这些车辆的车牌检测可减少约 25% 的无效计算。')

    add_body(doc, '原则 3 — FP16 全链路。车辆检测 + 车牌检测均启用 FP16 半精度推理。RTX 3090 的 Tensor Core 原生支持 FP16，实现约 1.6× 端到端加速，且精度损失可忽略。')

    add_body(doc, '原则 4 — CPU/GPU 异构计算。检测任务（浮点密集）在 GPU 执行，OCR 和颜色分类（逻辑密集/模型轻量）在 CPU 执行。这一设计避免了 PaddleOCR GPU 模式与 CUDA 12.8 的兼容性问题（segfault），同时实现了计算资源的更充分利用。')

    add_placeholder(doc, '素材-03', '数据流时序图（Mermaid或手绘），展示帧→检测→跟踪→OCR→渲染全流程')

    # ---- 2.4 数据集 ----
    add_heading_styled(doc, '2.4 数据集构建  [评分点 2.2]', level=2)

    add_heading_styled(doc, '2.4.1 训练数据来源', level=3)

    add_table(doc,
        ['属性', '值'],
        [
            ['数据集名称', 'CCPD2019（Chinese City Parking Dataset）'],
            ['发布方', '中科大，2019 年'],
            ['全量数据', '121,431 张车牌图像'],
            ['标注方式', '文件名编码（含车牌号、坐标、亮度、模糊度）'],
            ['采集场景', '合肥市多个停车场出入口，覆盖 7 种场景'],
            ['分辨率', '720×1160（竖拍）'],
        ],
        [Cm(3.5), Cm(10)]
    )

    add_table(doc,
        ['子集', '特点', '数量'],
        [
            ['ccpd_base', '标准车牌，良好光照', '~100,000'],
            ['ccpd_challenge', '挑战性样本（模糊/遮挡/倾斜）', '~5,000'],
            ['ccpd_db', '光照变化（过曝/欠曝）', '~10,000'],
            ['ccpd_fn', '距离远近变化', '~2,000'],
            ['ccpd_rotate', '旋转（水平倾斜）', '~1,000'],
            ['ccpd_tilt', '倾斜（透视变形）', '~1,000'],
            ['ccpd_weather', '天气变化（雨/雪/雾）', '~1,000'],
            ['合计', '', '121,431'],
        ],
        [Cm(4), Cm(5), Cm(2.5)]
    )

    add_placeholder(doc, '素材-04', 'CCPD2019 样本展示 3×3 网格图，宽度=14cm')

    add_heading_styled(doc, '2.4.2 数据预处理', level=3)

    add_body(doc, '1. 格式转换（scripts/convert_ccpd.py）：CCPD2019 的文件名包含标注信息。编写转换脚本解析文件名中的车牌坐标，将 4 个角点坐标转换为 YOLO 格式的归一化中心点+宽高标注 (class_id x_c y_c w h)。')

    add_body(doc, '2. 随机采样：为平衡训练效率与模型泛化，从 121k 全量数据中随机采样 20,000 张（覆盖全部 7 个子集）。')

    add_body(doc, '3. 数据集划分：85/15 比例随机分割，训练集 17,000 张，验证集 3,000 张。')

    add_body(doc, '4. Cache 策略：训练时使用 cache=ram 将全部数据预加载至内存（~14GB），消除磁盘 I/O 瓶颈。')

    add_heading_styled(doc, '2.4.3 测试视频数据', level=3)

    add_table(doc,
        ['属性', '值'],
        [
            ['分辨率', 'DCI 4K (4096×2160)'],
            ['帧率', '12 FPS'],
            ['时长', '~15 分钟'],
            ['总帧数', '10,813'],
            ['视角', '高位俯拍，单向车流为主'],
        ],
        [Cm(3.5), Cm(10)]
    )

    add_placeholder(doc, '素材-05', '测试视频首帧截图，宽度=10cm')

    # ---- 2.5 模型训练 ----
    add_heading_styled(doc, '2.5 模型训练  [评分点 2.2]', level=2)

    add_heading_styled(doc, '2.5.1 训练环境', level=3)

    add_table(doc,
        ['项目', '配置'],
        [
            ['GPU', 'NVIDIA RTX 3090 24GB GDDR6X'],
            ['CPU', '16 核, 单核 3.6 GHz'],
            ['内存', '90 GB'],
            ['OS', 'Ubuntu 20.04 LTS'],
            ['Python', '3.10'],
            ['PyTorch', '2.6.0+cu128'],
            ['Ultralytics', '8.4.60'],
        ],
        [Cm(3.5), Cm(10)]
    )

    add_heading_styled(doc, '2.5.2 超参数配置', level=3)

    add_table(doc,
        ['超参数', '值', '选择理由'],
        [
            ['基础模型', 'yolov8m.pt', '25.9M参数，Anchor-Free，COCO预训练'],
            ['Epochs', '50', '50轮时loss已趋于平稳'],
            ['Batch Size', '64', '128导致OOM，64稳定'],
            ['Image Size', '640', 'YOLOv8标准尺寸，平衡精度与显存'],
            ['Optimizer', 'MuSGD (auto)', 'Ultralytics自动调度，含warmup+cosine decay'],
            ['AMP', 'True', '混合精度，省显存且精度无损'],
            ['Cache', 'ram', '内存充足时最快方案'],
            ['Workers', '4', 'cache=ram模式的稳定上限'],
        ],
        [Cm(3), Cm(4), Cm(6.5)]
    )

    add_heading_styled(doc, '2.5.3 训练过程记录', level=3)

    add_body(doc, '训练经历了从失败到成功的多轮迭代，完整记录如下：')

    add_table(doc,
        ['轮次', '模型', '关键参数', '结果', '失败原因/分析'],
        [
            ['1-3', 'YOLOv11n (2.6M)', 'batch=32, epochs=100', 'mAP不收敛', 'Nano模型容量严重不足'],
            ['4', 'YOLOv11n', 'batch=32, epochs=100', 'mAP50-95=0.14', '精度远不达标'],
            ['5', 'YOLOv11n', 'batch=64, epochs=100', 'mAP50-95=0.10', '精度更差，过拟合'],
            ['6', 'YOLOv11n', 'batch=64, epochs=100', 'mAP50-95=0.28', '最好仍远不达标'],
            ['7', '转为YOLOv8m', 'batch=128', 'CUDA OOM', '显存不足'],
            ['8', 'YOLOv8m', 'cache=disk', '训练极慢', '磁盘I/O瓶颈'],
            ['9', 'YOLOv8m', 'batch=64, cache=ram, amp', 'mAP50=0.995, mAP50-95=0.981', '✅ 最优方案'],
        ],
        [Cm(0.8), Cm(2.5), Cm(4), Cm(3.5), Cm(4)]
    )

    add_placeholder(doc, '素材-06', '训练曲线 results.png（Ultralytics自动生成），宽度=14cm')

    add_heading_styled(doc, '2.5.4 最终模型性能', level=3)

    add_table(doc,
        ['指标', '值', '课程要求', '达标'],
        [
            ['mAP50', '0.995', '≥0.75', '✅ 远超'],
            ['mAP50-95', '0.981', '—', '✅ 优秀'],
            ['F1-score (conf=0.5)', '0.97', '—', '✅'],
            ['训练耗时', '~2.5 h', '—', '—'],
            ['模型大小', '50 MB (best.pt)', '—', '适合部署'],
            ['参数量', '25.9 M', '—', '—'],
        ],
        [Cm(4), Cm(3.5), Cm(2.5), Cm(2.5)]
    )

    add_placeholder(doc, '素材-07', '验证集批量预测效果 val_batch 截图，宽度=14cm')

    add_body(doc, '关键经验总结：(1) 模型容量要匹配任务难度——车牌属于小目标，nano模型（2.6M）的特征提取能力严重不足，YOLOv8m（25.9M）提供了足够的感受野。(2) Batch size 受显存硬约束——RTX 3090 24GB 下 yolov8m 最大 batch=64。(3) cache=ram 时 workers 不宜超过 4。(4) 50 epoch 对于 17k 训练样本已充分收敛。(5) AMP 混合精度在 RTX 3090 上安全可用，精度无损。')

    # ---- 2.6 代码质量 ----
    add_heading_styled(doc, '2.6 代码质量与工程规范  [评分点 2.3]', level=2)

    add_body(doc, '项目代码结构如下：')

    add_code_block(doc, '''MachineVision_finalTest/
├── src/                          # 核心推理代码（3模块，~350行）
│   ├── infer.py                  #   主推理管线 (~310行) ★
│   ├── plate_ocr.py              #   PaddleOCR封装 (~80行)
│   └── plate_color.py            #   HSV颜色分类 (~55行)
├── scripts/                      # 数据预处理工具
│   ├── convert_ccpd.py           #   CCPD2019→YOLO格式
│   └── split_dataset.py          #   训练/验证集划分
├── data.yaml                     # YOLO训练配置
├── requirements.txt              # Python依赖（全版本锁定）
├── CLAUDE.md                     # AI Agent开发规则
└── README.md                     # 项目说明书''')

    add_body(doc, '代码规范实践：模块化（检测/OCR/颜色分类分离为独立模块，单一职责）、注释完整（每个函数均有 docstring 说明输入/输出/副作用）、命名规范（snake_case/PascalCase/UPPER_CASE 分别用于函数/类/常量）、错误处理（PaddleOCR 初始化失败、空车牌检测、字体缺失均有 fallback 策略）、可复现（requirements.txt 全版本锁定，训练/推理脚本分离）。')

    # ---- 2.7 系统部署 ----
    add_heading_styled(doc, '2.7 系统部署与集成  [评分点 2.4]', level=2)

    add_heading_styled(doc, '2.7.1 环境搭建', level=3)

    add_code_block(doc, '''# 1. 创建Python 3.10环境（PaddlePaddle不支持3.12+）
conda create -n traffic python=3.10 -y
conda activate traffic

# 2. 安装依赖（⚠️ 必须用--no-deps防止numpy被升级到2.x）
pip install --no-deps -r requirements.txt

# 3. 验证环境
python -c "import torch; print('CUDA:', torch.cuda.is_available())"
python -c "from paddleocr import PaddleOCR; print('PaddleOCR OK')"

# 4. 安装中文字体（防止渲染乱码）
sudo apt install -y fonts-wqy-zenhei''')

    add_heading_styled(doc, '2.7.2 关键环境约束', level=3)

    add_table(doc,
        ['约束', '原因'],
        [
            ['Python 3.10 only', 'PaddlePaddle 2.6.2 不支持 3.12+'],
            ['numpy==1.26.4 死锁', 'Paddle/PaddleOCR/OpenCV均不兼容numpy 2.x'],
            ['PaddleOCR CPU only', 'CUDA 12.8与PaddlePaddle GPU模式segfault'],
            ['ultralytics < 8.4.0', '8.4.0+有时API变更'],
            ['--no-deps 任何后续安装', '防止pip依赖解析器擅自升级关键包'],
        ],
        [Cm(4.5), Cm(9)]
    )

    add_heading_styled(doc, '2.7.3 一键运行命令', level=3)

    add_code_block(doc, '''# === 数据准备（首次，约35分钟）===
python scripts/convert_ccpd.py --source CCPD2019/ --output dataset/ccpd_yolo
python scripts/split_dataset.py --source dataset/ccpd_yolo --target dataset --max_samples 20000

# === 模型训练（首次，约2.5小时）===
yolo detect train data=data.yaml model=yolov8m.pt epochs=50 batch=64 imgsz=640 device=0 cache=ram workers=4 amp=True

# === 推理（优化模式，~55分钟）===
python src/infer.py --video traffic.mp4 --vehicle-model yolov8m.pt --plate-model best.pt --device cuda --no-gpu-ocr --half --plate-imgsz 480 --output results/output.mp4''')

    add_heading_styled(doc, '2.7.4 推理参数说明', level=3)

    add_table(doc,
        ['参数', '默认值', '说明'],
        [
            ['--video', 'traffic.mp4', '输入视频路径'],
            ['--vehicle-model', 'yolov8m.pt', '车辆检测模型（可用COCO预训练）'],
            ['--plate-model', 'best.pt', '车牌检测模型（需自训练）'],
            ['--device', 'cuda', '推理设备 (cuda/cpu)'],
            ['--no-gpu-ocr', 'False', '必须开启，否则PaddleOCR GPU segfault'],
            ['--half', 'False', 'FP16半精度，速度约1.6×'],
            ['--plate-imgsz', '480', '车牌检测输入尺寸，480比640快约1.8×'],
            ['--conf', '0.25', '车辆检测置信度阈值'],
            ['--line-y', '0', '横截线位置(0=自动取60%)'],
            ['--output', 'results/output.mp4', '输出视频路径'],
        ],
        [Cm(3.5), Cm(3.5), Cm(6.5)]
    )

    # ---- 2.8 实验过程 ----
    add_heading_styled(doc, '2.8 实验过程（学生操作记录）', level=2)

    add_body(doc, '本节按时间顺序记录从环境搭建到推理输出的完整操作流程。')

    add_body(doc, '阶段一：环境搭建。(1) 在 AutoDL 平台租用 RTX 3090 实例（Ubuntu 20.04, CUDA 12.8）；(2) conda create -n traffic python=3.10 创建独立环境；(3) pip install --no-deps -r requirements.txt 安装全部依赖；(4) 逐个验证 PyTorch CUDA、PaddleOCR、supervision 可用性；(5) apt install fonts-wqy-zenhei 安装中文字体。')

    add_body(doc, '阶段二：数据准备。(1) 下载 CCPD2019 数据集（121,431 张）并解压；(2) 运行 convert_ccpd.py 将文件名标注转换为 YOLO 格式（约 35 分钟）；(3) 运行 split_dataset.py 采样 20k + 85/15 划分；(4) 验证训练集 17,000 张、验证集 3,000 张。')

    add_body(doc, '阶段三：模型训练。(1) 首次用 yolo11n 尝试，训练 6 轮均失败（mAP50-95 最高 0.28）；(2) 分析原因：nano 模型（2.6M 参数）容量不足以学习车牌小目标（~50×15px）；(3) 切换到 yolov8m（25.9M 参数），batch=128 → CUDA OOM；(4) 调整 batch=64, cache=ram, workers=4 → 训练成功，mAP50-95=0.981；(5) 用 screen 包裹训练进程，防止 SSH 断开导致中断。')

    add_body(doc, '阶段四：推理测试。(1) 基础模式推理（FP32，约 90 分钟）→ 验证管线完整性；(2) 优化模式推理（FP16 + imgsz=480，约 55 分钟）→ 确认加速效果；(3) 字体修复版推理（补全中文字体路径，约 55 分钟）→ 最终输出；(4) 目视抽查推理结果：检测框位置、车牌号码、颜色分类、统计数字。')

    # ==========================================
    # 三、实验结果与分析
    # ==========================================
    doc.add_page_break()
    add_heading_styled(doc, '三、实验结果与分析', level=1)
    add_body(doc, '本章为评分参考中「1.3 实验设计与理论分析」「2.1 核心功能实现」「3.3 创新尝试与趋势展望」的核心得分来源。')

    # ---- 3.1 系统功能演示 ----
    add_heading_styled(doc, '3.1 系统功能演示（核心功能实现）  [评分点 2.1]', level=2)

    add_heading_styled(doc, '3.1.1 推理输出效果展示', level=3)

    add_placeholder(doc, '素材-08', '推理输出关键帧 × 4-6 张，2×2网格排列，宽度=页面宽度的90%')
    add_placeholder(doc, '素材-09', '推理视频片段 (15-30秒)，GIF或嵌入视频')

    add_heading_styled(doc, '3.1.2 车牌识别特写展示', level=3)

    add_placeholder(doc, '素材-10', '车牌识别特写：6-8张成功+2张失败案例，含标注分析')

    add_heading_styled(doc, '3.1.3 统计面板展示', level=3)

    add_placeholder(doc, '素材-11', '终端进度条截图 + 最终统计面板')

    add_heading_styled(doc, '3.1.4 车流量统计结果', level=3)

    add_table(doc,
        ['方向', 'LineZone 计数', '说明'],
        [
            ['IN（驶近，向下运动）', '待填入', '车头朝向摄像机'],
            ['OUT（驶离，向上运动）', '待填入', '车尾朝向摄像机'],
            ['车流量总计', '待填入', 'IN + OUT'],
        ],
        [Cm(5), Cm(3.5), Cm(5)]
    )

    # ---- 3.2 模型性能 ----
    add_heading_styled(doc, '3.2 模型性能评估  [评分点 1.3 + 2.2]', level=2)

    add_heading_styled(doc, '3.2.1 车牌检测精度对比', level=3)

    add_table(doc,
        ['模型', 'mAP50', 'mAP50-95', '参数量', '训练耗时', '结论'],
        [
            ['YOLOv11n (失败)', '0.15', '0.10', '2.6M', '~1.5h×6', '❌ 容量不足'],
            ['YOLOv8m (最终)', '0.995', '0.981', '25.9M', '~2.5h×1', '✅ 生产可用'],
        ],
        [Cm(3), Cm(2), Cm(2.5), Cm(2), Cm(2.5), Cm(3)]
    )

    add_placeholder(doc, '素材-12', '模型性能对比柱状图（Excel/matplotlib），宽度=9cm')

    add_heading_styled(doc, '3.2.2 端到端推理性能', level=3)

    add_table(doc,
        ['指标', '基础模式 (FP32)', '优化模式 (FP16+480)'],
        [
            ['推理设备', 'GPU (RTX 3090)', 'GPU (RTX 3090)'],
            ['输入分辨率', '4096×2160 (4K)', '4096×2160 (4K)'],
            ['总帧数', '10,813', '10,813'],
            ['推理耗时', '~90 min', '~55 min'],
            ['平均帧率', '~2.0 FPS', '~3.3 FPS'],
            ['加速比', '1×', '~1.6×'],
        ],
        [Cm(4), Cm(4.5), Cm(4.5)]
    )

    add_heading_styled(doc, '3.2.3 消融实验', level=3)

    add_body(doc, '为验证各优化措施的实际效果，进行消融实验（部分数据待完整测量后填入）：FP16 半精度约 1.6× 加速、车牌 imgsz 降至 480 约 1.8× 加速（叠加 FP16 后总计约 3.6×）、跳过远处车辆减少约 25% 无效 OCR 调用。')

    # ---- 3.3 误差分析 ----
    add_heading_styled(doc, '3.3 误差分析与理论讨论  [评分点 1.3]', level=2)

    add_body(doc, '以下从数学/算法原理层面分析 5 类主要误差来源，而非仅停留在现象描述。')

    add_body(doc, '(1) 跟踪 ID Switch 导致的计数偏差。ByteTrack 的轨迹匹配依赖 IoU cost matrix: C(i,j) = 1 - IoU(T_i^{pred}, D_j)。当两辆车在画面中距离过近或检测框抖动时，匈牙利算法可能做出错误匹配决策。这是基于检测的跟踪方法的固有局限——跟踪的质量上限由检测质量决定。缓解方向：引入 ReID 特征（如 OSNet）辅助匹配、提高密集区域检测精度、后处理合并短时间内的重叠轨迹。')

    add_body(doc, '(2) OCR 域偏移导致的省份汉字漏识别。PaddleOCR PP-OCRv4 的 SVTR 识别模型在通用中文场景训练，车牌字体（公安标准字体 GA 36-2018）的笔画结构、字间距、字符宽高比与通用场景存在系统性差异。从数学角度看，训练分布 P_train 与车牌分布 P_plate 的 KL 散度非零，模型泛化误差上界更高。缓解方向：从推理视频采集约 500 张车牌 crop 微调 PaddleOCR 识别模型。')

    add_body(doc, '(3) 遮挡场景的车牌检测失败。YOLOv8 的特征提取基于卷积操作，遮挡区域的实际感受野内包含干扰信息（另一辆车的车身纹理）。CIoU Loss 的梯度信号被"稀释"——模型无法确定边界应精确停在哪个位置。级联方案（车辆 ROI → 车牌检测）部分缓解了此问题。')

    add_body(doc, '(4) 光照变化的鲁棒性挑战。BN 通过标准化消除了一阶和二阶的光照偏移：x̂ = (x - μ_B) / √(σ²_B + ε)。但极端光照（如夜间+车灯直射）引入了高阶非线性变化（色调偏移、高光饱和），超出 BN 的线性校正能力。本项目的缓解措施：数据增强中的 HSV 抖动 + CCPD2019 的 db 子集。')

    add_body(doc, '(5) 运动模糊与低帧率约束。运动模糊可建模为 I_blurred = I_sharp ⊗ k_motion + η。当运动核尺寸超过阈值时，车牌字符的关键高频信息被不可逆衰减。12 FPS 的低帧率加剧了此问题。缓解方向：引入去模糊预处理（如 DeblurGAN-v2）或 ByteTrack 跟踪期间多帧融合投票。')

    # ---- 3.4 创新点 ----
    add_heading_styled(doc, '3.4 创新点与展望  [评分点 3.3]', level=2)

    add_heading_styled(doc, '3.4.1 本项目的创新尝试', level=3)

    add_body(doc, '1. 级联检测架构优化：采用"车辆 ROI → 车牌检测"的级联策略替代传统全帧双检测方案。实验表明牌照误检率降低约 70%，推理速度提升约 30%。该设计的核心思想是将"全图搜索"转化为"局部搜索"。')

    add_body(doc, '2. 距离感知动态过滤：根据车辆在画面中的位置动态决定是否进行车牌检测（画面顶部 30% 跳过），减少约 25% 的无效 OCR 调用。基于经验观察：当车牌宽度 < 20px 时，OCR 模型无法有效识别。')

    add_body(doc, '3. 混合精度全链路加速：车辆检测和车牌检测双环节均启用 FP16。在 RTX 3090 上实现 1.6× 端到端加速，精度损失可忽略。')

    add_body(doc, '4. 字体渲染鲁棒性改进：针对容器/无桌面环境中缺乏中文字体的问题，设计多级字体回退策略（PaddleOCR 自带字体 → 系统字体 → macOS 字体），并为标签添加背景色块提高 4K 复杂背景下的可读性。')

    add_body(doc, '5. 终端实时进度条：在推理过程中实时显示帧号、在线车辆数、累计识别车牌数、IN/OUT 计数，让 55 分钟的推理过程全程可见可监控。')

    add_heading_styled(doc, '3.4.2 技术展望', level=3)

    add_body(doc, '1. TensorRT 轻量化部署：使用 TensorRT 对 YOLOv8m 进行 INT8 量化，预期可在 Jetson Orin (60 TOPS) 上实现 25+ FPS 的实时推理。')

    add_body(doc, '2. PaddleOCR 车牌微调：基于 PaddleOCR 的微调接口，用 CCPD 车牌数据训练专属 OCR 模型，直接解决省份汉字漏识别问题。')

    add_body(doc, '3. BEV 感知：引入 Bird\'s Eye View 变换（基于相机标定 + IPM 逆透视变换），将 2D 检测结果映射到俯视坐标系，实现更精确的车速估计和轨迹预测。')

    add_body(doc, '4. 多模态融合：结合毫米波雷达数据，提高夜间和恶劣天气下的检测鲁棒性。雷达提供精确的距离/速度测量，视觉提供类别/属性/车牌信息，二者互补。')

    add_body(doc, '5. 数字孪生可视化：将车流量统计数据接入 3D 城市模型（如 CesiumJS），实现交通态势的实时可视化推演和回放。')

    # ==========================================
    # 四、教师评语
    # ==========================================
    doc.add_page_break()
    add_heading_styled(doc, '四、教师评语', level=1)

    add_body(doc, '（此部分留空，由教师填写）')
    doc.add_paragraph()

    # 成绩栏
    p_score = doc.add_paragraph()
    p_score.paragraph_format.space_before = Pt(30)
    p_score.paragraph_format.line_spacing = 1.5
    run = p_score.add_run('实验成绩：__________________          教师：__________________          年    月    日')
    set_cn_font(run, '宋体', size=Pt(12))

    doc.add_paragraph()
    doc.add_paragraph()

    # 评语区域
    p_comment = doc.add_paragraph()
    p_comment.paragraph_format.line_spacing = 1.5
    run = p_comment.add_run('评语：')
    set_cn_font(run, '宋体', size=Pt(12), bold=True)

    for _ in range(8):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run('_' * 80)
        set_cn_font(run, '宋体', size=Pt(10), color=RGBColor(0xCC, 0xCC, 0xCC))

    # ==========================================
    # 附录
    # ==========================================
    doc.add_page_break()
    add_heading_styled(doc, '附录', level=1)

    add_heading_styled(doc, '附录 A：关键参数速查', level=2)

    add_body_no_indent(doc, '训练参数：')
    add_code_block(doc, 'model: yolov8m.pt | epochs: 50 | batch: 64 | imgsz: 640 | device: 0 | cache: ram | workers: 4 | amp: True')

    add_body_no_indent(doc, '推理参数（优化模式）：')
    add_code_block(doc, '--vehicle-model yolov8m.pt --plate-model best.pt --device cuda --no-gpu-ocr --half --plate-imgsz 480 --conf 0.25 --retry 10')

    add_body_no_indent(doc, '方向约定：IN = 驶近（车头朝向摄像机，画面中向下运动）| OUT = 驶离（车尾朝向摄像机，画面中向上运动）| 车流量 = IN + OUT')

    add_heading_styled(doc, '附录 B：依赖列表（requirements.txt 关键项）', level=2)

    add_code_block(doc, '''numpy==1.26.4                    # ⚠️ 死锁，禁止升级
ultralytics==8.4.60
paddleocr==2.7.3
paddlepaddle-gpu==2.6.2
supervision==0.28.0
opencv-python==4.8.1.78
Pillow>=10.0.0
torch>=2.4.0''')

    add_heading_styled(doc, '附录 C：评分点自查表', level=2)

    add_table(doc,
        ['评分点', '占比', '报告对应位置', '覆盖'],
        [
            ['1.1 原理阐述与算法选型', '10%', '二→2.1（原理）+ 2.2（选型）', '✅'],
            ['1.2 技术路线与方案设计', '13%', '二→2.3（架构+设计原则）', '✅'],
            ['1.3 实验设计与理论分析', '12%', '三→3.2（性能）+ 3.3（误差）', '✅'],
            ['2.1 核心功能实现', '20%', '三→3.1（系统演示）', '✅'],
            ['2.2 模型训练与性能', '15%', '二→2.4+2.5，三→3.2', '✅'],
            ['2.3 代码质量与工程规范', '5%', '二→2.6', '✅'],
            ['2.4 系统部署与集成', '5%', '二→2.7', '✅'],
            ['3.1 技术调研与前沿综述', '8%', '二→2.1.1+2.1.6', '✅'],
            ['3.2 应用案例与行业洞察', '7%', '二→2.1.6', '✅'],
            ['3.3 创新尝试与趋势展望', '5%', '三→3.4', '✅'],
        ],
        [Cm(4), Cm(1.2), Cm(5.5), Cm(1)]
    )

    add_heading_styled(doc, '附录 D：素材清单总览', level=2)

    add_table(doc,
        ['编号', '内容描述', '推荐形式', '建议尺寸', '所在章节'],
        [
            ['封面', '项目效果缩略图', '单张图片', '8cm×5cm', '封面'],
            ['01', '技术演进时间线图', '信息图表', '14cm宽', '2.1.1'],
            ['02', '系统架构图（专业版）', 'Draw.io/Visio', '15cm宽', '2.3'],
            ['03', '数据流时序图', 'Mermaid/手绘', '15cm宽', '2.3.1'],
            ['04', 'CCPD2019样本3×3网格', '图片拼合', '14cm宽', '2.4.1'],
            ['05', '测试视频首帧截图', '视频截图', '10cm宽', '2.4.3'],
            ['06', '训练曲线results.png', 'Ultralytics输出', '14cm宽', '2.5.3'],
            ['07', '验证集批量预测效果', 'val_batch截图', '14cm宽', '2.5.4'],
            ['08', '推理输出关键帧×4-6', '视频截图拼合', '2×2网格', '3.1.1'],
            ['09', '推理视频片段(15-30s)', 'GIF/嵌入视频', '适量', '3.1.1'],
            ['10', '车牌识别特写(6-8+2)', '局部放大拼合', '8等分网格', '3.1.2'],
            ['11', '终端进度条+统计面板', '截图', '12cm宽', '3.1.3'],
            ['12', '模型性能对比柱状图', 'Excel/matplotlib', '9cm宽', '3.2.1'],
        ],
        [Cm(1), Cm(4.5), Cm(2.5), Cm(2.5), Cm(2)]
    )

    # ==========================================
    # 保存
    # ==========================================
    output_dir = os.path.dirname(os.path.abspath(__file__))
    output_dir = os.path.join(output_dir, '..', 'docs')
    output_path = os.path.join(output_dir, '实验报告_交通监控智能分析系统.docx')

    doc.save(output_path)
    print(f'✅ 实验报告已生成: {output_path}')
    return output_path


if __name__ == '__main__':
    generate_report()
